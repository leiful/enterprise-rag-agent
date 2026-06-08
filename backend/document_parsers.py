# -*- coding: utf-8 -*-
"""Document parsers for extracting text from supported knowledge files."""

import re
from pathlib import Path
from typing import Optional, Tuple


def parse_document_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf_segments(file_path)
    if ext == ".docx":
        return parse_docx_segments(file_path)
    if ext in {".md", ".txt", ".rst"}:
        return parse_text_segments(file_path)
    if ext == ".csv":
        return parse_csv_segments(file_path)
    if ext in {".xlsx", ".xls"}:
        return parse_excel_segments(file_path)

    text, error = parse_document(file_path)
    if error:
        return [], error

    if not text.strip():
        return [], None

    return [
        {
            "text": text,
            "metadata": {},
        }
    ], None


def parse_text_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    text, error = parse_text_file(file_path)
    if error:
        return [], error
    if not text.strip():
        return [], None

    if file_path.suffix.lower() == ".md":
        segments = split_markdown_sections(text)
        if segments:
            return segments, None

    return [{"text": text, "metadata": {"section_path": ""}}], None


def split_markdown_sections(text):
    segments = []
    heading_stack = []
    current_lines = []

    def flush():
        body = "\n".join(current_lines).strip()
        if not body:
            return
        section_path = " > ".join(title for _, title in heading_stack)
        segments.append(
            {
                "text": body,
                "metadata": {
                    "section_path": section_path,
                    "section_title": heading_stack[-1][1] if heading_stack else "",
                },
            }
        )

    for line in text.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush()
            current_lines = [line]
            level = len(match.group(1))
            title = match.group(2).strip()
            heading_stack = [(item_level, item_title) for item_level, item_title in heading_stack if item_level < level]
            heading_stack.append((level, title))
            continue
        current_lines.append(line)

    flush()
    return segments


def parse_document(file_path: Path) -> Tuple[str, Optional[str]]:
    """Select a parser by file extension and return extracted text plus an error."""
    ext = file_path.suffix.lower()

    if ext in {".md", ".txt", ".rst"}:
        return parse_text_file(file_path)
    elif ext in {".html", ".htm"}:
        return parse_html_file(file_path)
    elif ext in {".docx"}:
        return parse_docx_file(file_path)
    elif ext in {".doc"}:
        return parse_doc_file(file_path)
    elif ext in {".pdf"}:
        return parse_pdf_file(file_path)
    elif ext in {".csv"}:
        return parse_csv_file(file_path)
    elif ext in {".xlsx", ".xls"}:
        return parse_excel_file(file_path)
    else:
        # Unknown extensions are treated as plain text as a best-effort fallback.
        return parse_text_file(file_path)



def parse_text_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse a plain text file."""
    try:
        text = file_path.read_text(encoding="utf-8")
        return text, None
    except UnicodeDecodeError:
        try:
            text = file_path.read_text(encoding="gbk")
            return text, None
        except Exception as error:
            return "", f"failed to read text file: {error}"
    except Exception as error:
        return "", f"failed to read text file: {error}"


def parse_html_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse an HTML file with simple tag cleanup."""
    try:
        text = file_path.read_text(encoding="utf-8")
        # Keep this lightweight; richer HTML normalization can be added later.
        text = re.sub(r'<script.*?>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style.*?>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text, None
    except UnicodeDecodeError:
        try:
            text = file_path.read_text(encoding="gbk")
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            return text, None
        except Exception as error:
            return "", f"failed to read html file: {error}"
    except Exception as error:
        return "", f"failed to parse html: {error}"


def parse_docx_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(cell_texts))
        return "\n".join(text_parts), None
    except ImportError:
        return "", "python-docx library not installed. Install with: pip install python-docx"
    except Exception as error:
        return "", f"failed to parse docx: {error}"


def parse_docx_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    try:
        from docx import Document
        doc = Document(file_path)
        segments = []
        page = 1
        page_parts = []
        paragraph_index = 0

        def flush_page():
            nonlocal page_parts
            text = "\n".join(part for part in page_parts if part).strip()
            if text:
                segments.append(
                    {
                        "text": text,
                        "metadata": {
                            "page_start": page,
                            "page_end": page,
                            "page_source": "docx_manual_break",
                        },
                    }
                )
            page_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text:
                paragraph_index += 1
                page_parts.append(paragraph.text)
            for run in paragraph.runs:
                if 'w:type="page"' in run._element.xml:
                    flush_page()
                    page += 1

        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text for cell in row.cells if cell.text]
                if cell_texts:
                    page_parts.append(" | ".join(cell_texts))

        flush_page()

        if not segments:
            return [], None
        return segments, None
    except ImportError:
        return [], "python-docx library not installed. Install with: pip install python-docx"
    except Exception as error:
        return [], f"failed to parse docx: {error}"


def parse_doc_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Return a clear error for legacy DOC files."""
    return "", "DOC file parsing requires extra libraries (antiword or pywin32). Try converting to DOCX first."


def parse_pdf_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse a PDF file into plain text."""
    try:
        import pypdf
        text_parts = []
        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            for page in pdf_reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    continue
        return "\n\n".join(text_parts), None
    except ImportError:
        return "", "pypdf library not installed. Install with: pip install pypdf"
    except Exception as error:
        return "", f"failed to parse pdf: {error}"


def parse_pdf_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    try:
        import pypdf
        segments = []
        page_count = 0
        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            for index, page in enumerate(pdf_reader.pages, start=1):
                page_count += 1
                try:
                    text = page.extract_text()
                    if text:
                        segments.append(
                            {
                                "text": text,
                                "metadata": {
                                    "page_start": index,
                                    "page_end": index,
                                    "page_source": "pdf",
                                },
                            }
                        )
                except Exception:
                    continue
        if page_count and not segments:
            return [], "no extractable PDF text found; this may be a scanned PDF and requires OCR before indexing"
        return segments, None
    except ImportError:
        return [], "pypdf library not installed. Install with: pip install pypdf"
    except Exception as error:
        return [], f"failed to parse pdf: {error}"


def parse_csv_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse a CSV file into row-oriented text."""
    try:
        import csv
        text_parts = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(" | ".join(row))
        return "\n".join(text_parts), None
    except Exception as error:
        try:
            with open(file_path, "r", encoding="gbk") as f:
                reader = csv.reader(f)
                text_parts = []
                for row in reader:
                    text_parts.append(" | ".join(row))
                return "\n".join(text_parts), None
        except Exception as e:
            return "", f"failed to parse csv: {error}"


def parse_csv_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    try:
        import csv
        with open(file_path, "r", encoding="utf-8") as f:
            return csv_rows_to_segments(csv.reader(f), file_path.stem), None
    except UnicodeDecodeError:
        try:
            import csv
            with open(file_path, "r", encoding="gbk") as f:
                return csv_rows_to_segments(csv.reader(f), file_path.stem), None
        except Exception as error:
            return [], f"failed to parse csv: {error}"
    except Exception as error:
        return [], f"failed to parse csv: {error}"


def csv_rows_to_segments(rows, sheet_name):
    rows = [list(row) for row in rows]
    if not rows:
        return []

    header = [str(value).strip() for value in rows[0]]
    segments = []
    for row_index, row in enumerate(rows[1:], start=2):
        values = [str(value).strip() for value in row]
        if not any(values):
            continue
        pairs = []
        for cell_index, value in enumerate(values):
            column = header[cell_index] if cell_index < len(header) and header[cell_index] else f"Column {cell_index + 1}"
            pairs.append(f"{column}: {value}")
        searchable_row = " | ".join(pairs)
        segments.append(
            {
                "text": (
                    f"Sheet: {sheet_name}\n"
                    f"Row {row_index}\n"
                    f"Columns: {', '.join(header)}\n"
                    f"Record: {searchable_row}\n"
                    + "\n".join(pairs)
                ),
                "metadata": {
                    "sheet_name": sheet_name,
                    "row_start": row_index,
                    "row_end": row_index,
                    "table_header": header,
                },
            }
        )
    return segments


def parse_excel_file(file_path: Path) -> Tuple[str, Optional[str]]:
    """Parse an Excel workbook into sheet-oriented text."""
    try:
        import openpyxl
        text_parts = []
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        for sheet_name in wb.sheetnames:
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            sheet = wb[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                cell_texts = []
                for cell_value in row:
                    if cell_value is not None:
                        cell_texts.append(str(cell_value))
                if cell_texts:
                    text_parts.append(" | ".join(cell_texts))
        return "\n".join(text_parts), None
    except ImportError:
        return "", "openpyxl library not installed. Install with: pip install openpyxl"
    except Exception as error:
        return "", f"failed to parse excel: {error}"


def parse_excel_segments(file_path: Path) -> Tuple[list[dict], Optional[str]]:
    try:
        import openpyxl
        segments = []
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = [
                ["" if value is None else str(value) for value in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            segments.extend(csv_rows_to_segments(rows, sheet_name))
        return segments, None
    except ImportError:
        return [], "openpyxl library not installed. Install with: pip install openpyxl"
    except Exception as error:
        return [], f"failed to parse excel: {error}"
